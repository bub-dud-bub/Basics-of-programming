import PySimpleGUI as sg
from docx import Document
import openpyxl
import numpy as np
from cars.Cars import Mashina, Truck, Bus
class Super:
    def load(self):
        form = sg.FlexForm('Simple data entry form')
        layout = [[
              sg.Combo(['Легковушка', 'Грузовушка', 'Пассажирский'], key='type', size=(30, 6))],
              [sg.Text("Масса вместе с грузом(кг):", size=(20, 1)), sg.InputText('1000', key='mas')],
              [sg.Text("Путь(км):", size=(10, 1)), sg.InputText('1000', key='way')],
              [sg.Submit("Рассчитать")]
             ]
        window = sg.Window("Demo", layout)

        while True:
            event, values = window.read()
            self.cb = values['type']
            self.mas = values['mas']
            self.way = values['way']
            if event == sg.WIN_CLOSED:
                break
            if event == "Рассчитать":
                self.Results()

    def Results(self):
        if self.cb == 'Легковушка':
            fuel = Mashina.Fuel_sp(mas=self.mas)
            time = Mashina.Time(way=self.way)
            cost = Mashina.Cost(fuel=fuel, way=self.way)
        elif self.cb == 'Пассажирский':
            fuel = Bus.Fuel_sp(mas=self.mas)
            time = Bus.Time(way=self.way)
            cost = Bus.Cost(fuel=fuel, way=self.way)
        elif self.cb == 'Грузовушка':
            fuel = Truck.Fuel_sp(mas=self.mas)
            time = Truck.Time(way=self.way)
            cost = Truck.Cost(fuel=fuel, way=self.way)
        layout = [
              [sg.Text("Расход топлива: "+str(fuel) + " л. на 1 км.", size=(30, 1))],
              [sg.Text("Время поездки: "+str(time) + " часов", size=(30, 1))],
              [sg.Text("Стоимость: "+str(cost) + " руб.", size=(30, 1))],
              [sg.Submit("Сохранить в docx как:", size=(15, 1)), sg.InputText('Отчёрт', key='docx')],
              [sg.Submit("Сохранить в xlsx как:", size=(15, 1)), sg.InputText('Отчёрт', key='xlsx')]
             ]
        count = sg.Window("Demo", layout)
        self.fuel=fuel
        self.time=time
        self.cost=cost

        while True:
            event, values = count.read()
            if event == sg.WIN_CLOSED:
                break
            if event == "Сохранить в docx как:":
                self.sav = values['docx']
                self.Docx()
            if event == "Сохранить в xlsx как:":
                self.savv = values['xlsx']
                self.Xlsx()


    def Docx(self):
        doc = Document()
        head = doc.add_heading('Типа отчёт')
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        head.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p0 = doc.add_paragraph('Тип автомобиля: ' + self.cb)
        p = doc.add_paragraph('Масса вместе с грузом(кг): ' + self.mas)
        p2 = doc.add_paragraph('Путь(км): ' + self.way)
        p3 = doc.add_paragraph("Расход топлива:" + str(self.fuel) + " л. на 1 км.")
        p4 = doc.add_paragraph("Время поездки:" + str(self.time) + " часов")
        p5 = doc.add_paragraph("Стоимость:" + str(self.cost) + " руб.")
        doc.save(self.sav + '.docx')

    def Xlsx(self):
        xl = openpyxl.Workbook()
        sheet = xl.active
        sheet['B1'] = "Отчёт типа"
        sheet['A2'] = 'Тип автомобиля: ' + self.cb
        sheet['A3'] = 'Масса вместе с грузом(кг): '
        sheet['A4'] = 'Путь(км): ' + self.way
        sheet['A5'] = "Расход топлива:"
        sheet['A6'] = "Время поездки:"
        sheet['A7'] = "Стоимость:"

        sheet['B2'] = self.cb
        sheet['B3'] = self.mas
        sheet['B4'] = self.way
        sheet['B5'] = str(self.fuel) + " л. на 1 км."
        sheet['B6'] = str(self.time) + " часов"
        sheet['A7'] = str(self.cost) + " руб."
        xl.save(self.savv + '.xlsx')

class Inheritor(Super):
    pass

a = Inheritor()
a.load()
