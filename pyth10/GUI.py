import wx
from docx import Document
import openpyxl
class MyFrame(wx.Frame):
        def __init__(self, parent, title):
            wx.Frame.__init__(self, parent, title = title, size = (300,250))

            self.lbl1 = wx.StaticText(self, wx.ID_ANY, "Тип автомобиля:", pos=(10, 12))
            self.links = ["Легковой", "Грузовой", "Пассажирский"]
            self.cb = wx.ComboBox(self, pos=(110, 6), choices=self.links, style=wx.CB_READONLY)
            self.Show(True)

            self.lbl2 = wx.StaticText(self, wx.ID_ANY, "Масса вместе с грузом(кг):", pos=(10, 32))
            self.mas = wx.TextCtrl(self, wx.ID_ANY, "1000", pos=(165, 29))

            self.lbl3 = wx.StaticText(self, wx.ID_ANY, "Путь(км):", pos=(10, 52))
            self.way = wx.TextCtrl(self, wx.ID_ANY, "1000", pos=(70, 50))

            bt = wx.Button(self, wx.ID_ANY, "Рассчитать", pos=(10, 75))
            bt.Bind(wx.EVT_BUTTON, self.OnButton)

        def OnButton(self, event):
            if self.cb.Value == self.links[0]:
                import cars.Mashina as msh
            elif self.cb.Value == self.links[2]:
                import cars.Marshrutka as msh
            elif self.cb.Value == self.links[1]:
                import cars.Truck as msh
            self.fuel = msh.Fuel_sp(self.way)
            self.time = msh.Time(self.way.Value)
            self.cost = msh.Cost(self.mas.Value, self.fuel)
            self.lbl4 = wx.StaticText(self, wx.ID_ANY, "Расход топлива: "+str(self.fuel) + " л. на 1 км.", pos=(10, 100))
            self.lbl5 = wx.StaticText(self, wx.ID_ANY, "Время поездки: "+str(self.time) + " часов", pos=(10, 125))
            self.lbl6 = wx.StaticText(self, wx.ID_ANY, "Стоимость: "+str(self.cost) + " руб.", pos=(10, 150))
            bt2 = wx.Button(self, wx.ID_ANY, "Сохранить в docx как:", pos=(10, 175))
            self.sav = wx.TextCtrl(self, wx.ID_ANY, "Отчёт", pos=(150, 175))
            bt2.Bind(wx.EVT_BUTTON, self.OnButton2)
            bt3 = wx.Button(self, wx.ID_ANY, "Сохранить в xlsx как:", pos=(10, 200))
            self.savv = wx.TextCtrl(self, wx.ID_ANY, "Отчёт", pos=(150, 200))
            bt3.Bind(wx.EVT_BUTTON, self.OnButton3)

        def OnButton2(self, event):
            doc = Document()
            head = doc.add_heading('Типа отчёт')
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            head.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p0 = doc.add_paragraph('Тип автомобиля: ' + self.cb.Value)
            p = doc.add_paragraph('Масса вместе с грузом(кг): ' + self.mas.Value)
            p2 = doc.add_paragraph('Путь(км): ' + self.way.Value)
            p3 = doc.add_paragraph(self.lbl4.Label)
            p4 = doc.add_paragraph(self.lbl5.Label)
            p5 = doc.add_paragraph(self.lbl6.Label)
            doc.save(self.sav.Value + '.docx')
        def OnButton3(self, event):
            xl = openpyxl.Workbook()
            sheet = xl.active
            sheet['B1'] = "Отчёт типа"
            sheet['A2'] = self.lbl1.Label
            sheet['A3'] = self.lbl2.Label
            sheet['A4'] = self.lbl3.Label
            sheet['A5'] = "Расход топлива:"
            sheet['A6'] = "Время поездки:"
            sheet['A7'] = "Стоимость:"

            sheet['B2'] = self.cb.Value
            sheet['B3'] = self.mas.Value
            sheet['B4'] = self.way.Value
            sheet['B5'] = str(self.fuel) + " л. на 1 км."
            sheet['B6'] = str(self.time) + " часов"
            sheet['A7'] = str(self.cost) + " руб."
            xl.save(self.savv.Value + '.xlsx')
app = wx.App()
frame = MyFrame(None, "Автомобили")
app.MainLoop()
